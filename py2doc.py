#!/usr/bin/env python
"""
Requires:
  pip install python-docx pygments
"""

import os
import sys

# Dependency checks with clear guidance
try:
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Missing dependency: python-docx\nInstall with:\n  python -m pip install python-docx")
    sys.exit(1)

try:
    from pygments import highlight
    from pygments.lexers import PythonLexer
    from pygments.token import Token
    from pygments.formatter import Formatter
    HAVE_PYGMENTS = True
except ImportError:
    HAVE_PYGMENTS = False


def set_run_consolas(run):
    run.font.name = "Consolas"
    # Ensure Word respects the font for all scripts
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Consolas')


def set_paragraph_consolas(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    # No extra spacing between lines
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), "0")
    spacing.set(qn('w:after'), "0")
    pPr.append(spacing)


class DocxFormatter(Formatter):
    """
    Pygments -> python-docx formatter:
    - Consolas 10pt
    - Line numbers
    - Basic syntax coloring
    """
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.paragraph = None
        self.line_no = 0

    def new_line(self):
        self.paragraph = self.doc.add_paragraph()
        set_paragraph_consolas(self.paragraph)
        self.line_no += 1
        ln_run = self.paragraph.add_run(f"{self.line_no:>4} ")
        set_run_consolas(ln_run)
        ln_run.bold = True
        ln_run.font.size = Pt(10)

    def format(self, tokensource, outfile):
        self.new_line()
        for ttype, value in tokensource:
            # Handle newlines inside a token
            parts = value.split('\n')
            for i, part in enumerate(parts):
                if part:
                    run = self.paragraph.add_run(part)
                    set_run_consolas(run)
                    run.font.size = Pt(10)
                    self.apply_style(run, ttype)
                if i < len(parts) - 1:
                    self.new_line()

    def apply_style(self, run, ttype):
        # Colors and styles
        if ttype in Token.Keyword:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            run.bold = True
        elif ttype in Token.String:
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        elif ttype in Token.Comment:
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            run.italic = True
        elif ttype in Token.Name.Function:
            run.font.color.rgb = RGBColor(0xB0, 0x00, 0xB0)
        elif ttype in Token.Number:
            run.font.color.rgb = RGBColor(0x1A, 0x7F, 0x37)
        else:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def convert_py_to_docx(py_file_path, out_path=None, add_title=True):
    # Read source with a sensible fallback
    try:
        code = open(py_file_path, 'r', encoding='utf-8').read()
    except UnicodeDecodeError:
        code = open(py_file_path, 'r', encoding='cp1252', errors='replace').read()

    doc = Document()

    if add_title:
        doc.add_heading(f"Python Script: {os.path.basename(py_file_path)}", level=1)

    if HAVE_PYGMENTS:
        formatter = DocxFormatter(doc)
        highlight(code, PythonLexer(), formatter)
    else:
        # Plain fallback with line numbers and Consolas
        lines = code.splitlines()
        for i, line in enumerate(lines, start=1):
            p = doc.add_paragraph()
            set_paragraph_consolas(p)
            ln = p.add_run(f"{i:>4} ")
            set_run_consolas(ln)
            ln.bold = True
            ln.font.size = Pt(10)

            rr = p.add_run(line)
            set_run_consolas(rr)
            rr.font.size = Pt(10)

    if not out_path:
        base, _ = os.path.splitext(py_file_path)
        out_path = f"{base}.docx"

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 2:
        print("Usage:\n  python py2doc.py path/to/script.py [output.docx]\n\nIf you see 'Missing dependency', install with:\n  python -m pip install python-docx pygments")
        sys.exit(0)

    in_path = sys.argv[1]
    out_path = sys.argv[2] if len(sys.argv) >= 3 else None

    if not os.path.isfile(in_path):
        print(f"Not found: {in_path}")
        sys.exit(1)

    try:
        out_file = convert_py_to_docx(in_path, out_path)
    except Exception as e:
        print(f"Conversion failed: {e}")
        sys.exit(1)

    print(f"Saved: {out_file}")


if __name__ == "__main__":
    main()
